import fitz  # PyMuPDF
import re
import os
from pathlib import Path


def extract_text_from_pdf(pdf_path: str) -> dict:
    """
    Extract text from digital PDF using PyMuPDF
    Does NOT support scanned PDFs — use manual input instead
    """
    doc = fitz.open(pdf_path)
    full_text = ""
    pages_text = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text").strip()
        pages_text.append({
            "page_number": page_num + 1,
            "text": page_text
        })
        full_text += page_text + "\n"

    doc.close()

    # Check if PDF is scanned (very little text extracted)
    if len(full_text.strip()) < 100:
        return {
            "full_text": "",
            "pages": pages_text,
            "total_pages": len(pages_text),
            "total_chars": 0,
            "is_scanned": True,
            "warning": "This appears to be a scanned PDF. Please use manual unit input instead."
        }

    return {
        "full_text": full_text,
        "pages": pages_text,
        "total_pages": len(pages_text),
        "total_chars": len(full_text),
        "is_scanned": False
    }


def extract_text_from_docx(file_path: str) -> dict:
    """Extract text from Word document"""
    from docx import Document
    doc = Document(file_path)
    full_text = ""
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text + "\n"
            paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text += cell.text + "\n"

    return {
        "full_text": full_text,
        "pages": [{"page_number": 1, "text": full_text}],
        "total_pages": len(paragraphs),
        "total_chars": len(full_text),
        "is_scanned": False
    }


def extract_text_from_image(file_path: str) -> dict:
    """Extract text from image using Tesseract OCR"""
    import pytesseract
    from PIL import Image

    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)

    return {
        "full_text": text,
        "pages": [{"page_number": 1, "text": text}],
        "total_pages": 1,
        "total_chars": len(text),
        "is_scanned": False
    }


def extract_text(file_path: str) -> dict:
    """
    MAIN FUNCTION: Auto-detect file type and extract text
    Supports: PDF, DOCX, DOC, PNG, JPG, JPEG
    """
    ext = Path(file_path).suffix.lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext in ['.png', '.jpg', '.jpeg']:
        return extract_text_from_image(file_path)
    else:
        return {
            "full_text": "",
            "pages": [],
            "total_pages": 0,
            "total_chars": 0,
            "error": f"Unsupported file format: {ext}"
        }


def clean_text(text: str) -> str:
    """
    Advanced text cleaning for NLP processing
    Removes formatting artifacts from PDFs
    """
    # Remove excessive whitespace and special chars
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\r', '\n', text)

    # Remove bullet points and special markers
    text = re.sub(r'[●•·▪▸►◆★✓✔☐☑]', '', text)

    # Remove multiple spaces
    text = re.sub(r' {2,}', ' ', text)

    # Remove lines that are just symbols or numbers
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        # Skip empty lines
        if not line:
            continue
        # Skip very short lines (page numbers, headers)
        if len(line) < 15:
            continue
        # Skip lines that are mostly symbols
        alpha_ratio = sum(c.isalpha() for c in line) / max(len(line), 1)
        if alpha_ratio < 0.4:
            continue
        # Clean remaining formatting artifacts
        line = re.sub(r'\s+', ' ', line)
        line = re.sub(r'_{2,}', '', line)
        line = re.sub(r'-{2,}', '', line)
        line = re.sub(r'\|', ' ', line)
        # Skip memory hooks, tips, and exam notes
        skip_phrases = [
            'memory hook', 'watch out', 'key distinction',
            'tip:', 'note:', 'remember:', 'hint:',
            'exam tip', 'mark scheme', 'confidential info always',
            'doac', 'khap'
        ]
        if any(phrase in line.lower() for phrase in skip_phrases):
            continue
        cleaned_lines.append(line)

    # Join lines into proper sentences
    result = ' '.join(cleaned_lines)

    # Fix spacing around punctuation
    result = re.sub(r'\s+([.,;:!?])', r'\1', result)
    result = re.sub(r'([.,;:!?])\s*', r'\1 ', result)

    # Remove duplicate spaces again
    result = re.sub(r' {2,}', ' ', result)

    return result.strip()


def extract_sentences(text: str) -> list:
    """Split text into clean sentences"""
    import spacy
    nlp = spacy.load("en_core_web_lg")
    doc = nlp(text[:100000])
    sentences = [
        sent.text.strip()
        for sent in doc.sents
        if len(sent.text.strip()) > 20
    ]
    return sentences